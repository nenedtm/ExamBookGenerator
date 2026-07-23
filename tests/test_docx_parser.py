"""Tests for parsers.docx_parser."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches
from PIL import Image

from core.models import Document, FileType
from parsers.docx_parser import parse_docx


# ── Helpers ─────────────────────────────────────────────────────────────────

def _create_docx(
    path: Path,
    paragraphs: list[str] | None = None,
    title: str = "",
    add_image: bool = False,
) -> Path:
    """Build a minimal DOCX for testing."""
    doc = DocxDocument()
    if title:
        doc.core_properties.title = title

    for text in (paragraphs if paragraphs is not None else ["Hello world"]):
        doc.add_paragraph(text)

    if add_image:
        img = Image.new("RGB", (200, 150), color="blue")
        buf = BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        doc.add_picture(buf, width=Inches(2))

    doc.save(path)
    return path


def _create_docx_with_table(path: Path) -> Path:
    doc = DocxDocument()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "D"
    table.cell(1, 1).text = "E"
    table.cell(1, 2).text = "F"
    doc.add_paragraph("After table")
    doc.save(path)
    return path


def _corrupt_docx(path: Path) -> Path:
    path.write_bytes(b"PK\x03\x04corrupt zip data here")
    return path


# ── parse_docx ──────────────────────────────────────────────────────────────

class TestParseDocx:
    def test_returns_document(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx")
        doc, images = parse_docx(p)
        assert isinstance(doc, Document)
        assert isinstance(images, list)

    def test_file_type_is_docx(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx")
        doc, _ = parse_docx(p)
        assert doc.file_type == FileType.DOCX

    def test_source_path_matches(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx")
        doc, _ = parse_docx(p)
        assert doc.source_path == str(p)

    def test_text_extracted(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", paragraphs=["Alpha beta gamma"])
        doc, _ = parse_docx(p)
        assert "Alpha beta gamma" in doc.content

    def test_multiple_paragraphs(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", paragraphs=["First", "Second", "Third"])
        doc, _ = parse_docx(p)
        assert "First" in doc.content
        assert "Second" in doc.content
        assert "Third" in doc.content

    def test_table_text(self, tmp_path: Path) -> None:
        p = _create_docx_with_table(tmp_path / "table.docx")
        doc, _ = parse_docx(p)
        assert "A" in doc.content
        assert "B" in doc.content
        assert "Before table" in doc.content
        assert "After table" in doc.content

    def test_metadata_title(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", title="My Document")
        doc, _ = parse_docx(p)
        assert doc.title == "My Document"
        assert doc.metadata.get("title") == "My Document"

    def test_images_extracted(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", add_image=True)
        doc, images = parse_docx(p)
        assert len(images) >= 1
        assert len(doc.images) >= 1

    def test_image_file_on_disk(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", add_image=True)
        _, images = parse_docx(p)
        assert Path(images[0].file_path).exists()

    def test_no_images_doc(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", add_image=False)
        doc, images = parse_docx(p)
        assert images == []
        assert doc.images == []

    def test_empty_docx(self, tmp_path: Path) -> None:
        p = _create_docx(tmp_path / "test.docx", paragraphs=[])
        doc, images = parse_docx(p)
        assert doc.content.strip() == ""
        assert images == []

    def test_corrupt_docx(self, tmp_path: Path) -> None:
        _corrupt_docx(tmp_path / "bad.docx")
        doc, images = parse_docx(tmp_path / "bad.docx")
        assert doc.content == ""
        assert images == []

    def test_nonexistent_raises(self) -> None:
        with pytest.raises(FileNotFoundError):
            parse_docx(Path("/no/such/file.docx"))
